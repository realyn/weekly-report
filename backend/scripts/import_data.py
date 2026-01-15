#!/usr/bin/env python3
"""
从 Word 文档导入周报数据

用法:
    python scripts/import_data.py /path/to/docx/files/

该脚本会:
1. 解析 Word 文档中的周报表格
2. 创建用户账号（如不存在）
3. 导入周报数据
"""

import asyncio
import re
import sys
import os
from pathlib import Path
from datetime import datetime
from docx import Document

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from sqlalchemy import select
from app.database import init_db, async_session
from app.models.user import User, UserRole
from app.models.report import Report, ReportStatus
from app.utils.security import get_password_hash


# 员工信息配置
EMPLOYEES = {
    '杨宁': {'department': '产品研发部'},
    '翦磊': {'department': '产品研发部'},
    '朱迪': {'department': '产品研发部'},
    '张士健': {'department': '产品研发部'},
    '程志强': {'department': '产品研发部'},
    '闻世坤': {'department': '产品研发部'},
    '秦闪闪': {'department': '产品研发部'},
    '蒋奇朴': {'department': '产品研发部'},
}

# 默认用户密码
DEFAULT_PASSWORD = "Weekly@2026"


def parse_weekly_report(filepath: str) -> dict:
    """解析周报 Word 文档"""
    doc = Document(filepath)

    if not doc.tables:
        raise ValueError(f"文档 {filepath} 中没有找到表格")

    table = doc.tables[0]

    # 提取日期和周次
    header_text = ""
    for row in table.rows[:3]:
        for cell in row.cells:
            text = cell.text.strip()
            if "周工作周报" in text:
                header_text = text
                break
        if header_text:
            break

    # 解析日期
    date_match = re.search(r'(\d{4})年(\d{2})月(\d{2})日', header_text)
    week_match = re.search(r'第(\d+)周', header_text)

    if not date_match:
        raise ValueError(f"无法从 {filepath} 中解析日期")

    year, month, day = date_match.groups()

    # 优先从文件名提取年份（文件名格式如：产品研发部工作计划2026.01.10.docx）
    filename = os.path.basename(filepath)
    filename_year_match = re.search(r'(\d{4})\.(\d{2})\.(\d{2})', filename)
    if filename_year_match:
        year = filename_year_match.group(1)

    report_date = datetime.strptime(f"{year}-{month}-{day}", "%Y-%m-%d").date()
    week_number = int(week_match.group(1)) if week_match else 1

    # 找到"下周"部分的起始行
    next_week_start_row = None
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            text = cell.text.strip()
            # 移除空格后检查（文档中可能有空格："下    周   \n内    容    说    明"）
            text_no_space = text.replace(' ', '').replace('\n', '')
            if '下周' in text_no_space and ('内容' in text_no_space or '说明' in text_no_space):
                next_week_start_row = row_idx
                break
        if next_week_start_row is not None:
            break

    # 解析每个人的工作内容 - 逐行处理
    work_items = {}
    employee_names = list(EMPLOYEES.keys())

    for row_idx, row in enumerate(table.rows):
        # 获取该行所有单元格的文本（去重，保持顺序）
        row_cells = []
        seen_in_row = set()
        for cell in row.cells:
            text = cell.text.strip()
            if text and text not in seen_in_row:
                seen_in_row.add(text)
                row_cells.append(text)

        # 在该行中查找员工名
        for i, cell_text in enumerate(row_cells):
            if cell_text in employee_names:
                employee_name = cell_text
                if employee_name not in work_items:
                    work_items[employee_name] = {'this_week': [], 'next_week': []}

                # 在同一行中查找工作内容（在员工名前面的单元格中）
                for j in range(i):
                    content = row_cells[j]
                    # 检查是否是工作内容（以数字开头，且包含点或顿号）
                    if content and len(content) > 3 and content[0].isdigit():
                        if '.' in content[:4] or '、' in content[:4] or '，' in content[:4]:
                            # 根据行号判断是本周还是下周
                            is_next_week = next_week_start_row is not None and row_idx >= next_week_start_row

                            if is_next_week:
                                if content not in work_items[employee_name]['next_week']:
                                    work_items[employee_name]['next_week'].append(content)
                            else:
                                if content not in work_items[employee_name]['this_week']:
                                    work_items[employee_name]['this_week'].append(content)

    return {
        'date': report_date,
        'week': week_number,
        'year': int(year),
        'items': work_items
    }


async def import_data(docx_dir: str, update_existing: bool = False):
    """导入数据到数据库

    Args:
        docx_dir: Word文档目录
        update_existing: 是否更新已存在的记录
    """
    await init_db()

    # 查找所有 docx 文件
    docx_files = list(Path(docx_dir).glob("*.docx"))
    if not docx_files:
        print(f"在 {docx_dir} 中没有找到 .docx 文件")
        return

    print(f"找到 {len(docx_files)} 个文档")

    async with async_session() as db:
        # 1. 创建用户
        print("\n=== 创建用户 ===")
        users = {}
        for name, info in EMPLOYEES.items():
            result = await db.execute(select(User).where(User.real_name == name))
            user = result.scalar_one_or_none()

            if not user:
                # 生成用户名（拼音首字母）
                username = name.lower().replace(' ', '')
                # 检查用户名是否存在
                result = await db.execute(select(User).where(User.username == username))
                if result.scalar_one_or_none():
                    username = f"{username}_{len(users)}"

                user = User(
                    username=username,
                    password=get_password_hash(DEFAULT_PASSWORD),
                    real_name=name,
                    department=info['department'],
                    role=UserRole.user
                )
                db.add(user)
                print(f"  创建用户: {name} ({username})")
            else:
                print(f"  用户已存在: {name}")

            users[name] = user

        await db.commit()

        # 刷新用户对象以获取 ID
        for name in users:
            await db.refresh(users[name])

        # 2. 导入周报数据
        print("\n=== 导入周报 ===")
        for docx_file in sorted(docx_files):
            print(f"\n处理: {docx_file.name}")
            try:
                data = parse_weekly_report(str(docx_file))
                print(f"  日期: {data['date']}, 第{data['week']}周")

                for name, work in data['items'].items():
                    if name not in users:
                        continue

                    user = users[name]

                    # 检查是否已存在
                    result = await db.execute(
                        select(Report).where(
                            Report.user_id == user.id,
                            Report.year == data['year'],
                            Report.week_num == data['week']
                        )
                    )
                    existing = result.scalar_one_or_none()

                    this_week = '\n'.join(work.get('this_week', []))
                    next_week = '\n'.join(work.get('next_week', []))

                    if existing:
                        if update_existing:
                            existing.this_week_work = this_week or existing.this_week_work
                            existing.next_week_plan = next_week or existing.next_week_plan
                            print(f"  {name}: 更新成功")
                        else:
                            print(f"  {name}: 周报已存在，跳过")
                        continue

                    if this_week or next_week:
                        report = Report(
                            user_id=user.id,
                            year=data['year'],
                            week_num=data['week'],
                            this_week_work=this_week or "无",
                            next_week_plan=next_week or "待定",
                            status=ReportStatus.submitted
                        )
                        db.add(report)
                        print(f"  {name}: 导入成功 (本周{len(work.get('this_week', []))}条, 下周{len(work.get('next_week', []))}条)")

                await db.commit()

            except Exception as e:
                print(f"  解析错误: {e}")
                import traceback
                traceback.print_exc()
                continue

    print("\n=== 导入完成 ===")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python scripts/import_data.py /path/to/docx/files/ [--update]")
        print("  --update: 更新已存在的记录")
        sys.exit(1)

    docx_dir = sys.argv[1]
    update_existing = "--update" in sys.argv

    if not os.path.isdir(docx_dir):
        print(f"目录不存在: {docx_dir}")
        sys.exit(1)

    asyncio.run(import_data(docx_dir, update_existing))
