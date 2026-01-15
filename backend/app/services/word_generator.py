from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime, timedelta
from app.config import get_settings

settings = get_settings()


def set_row_height(row, height_pt):
    """设置行高（磅）"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))  # 转换为 twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def set_table_cell_margins(table, top_pt=0, bottom_pt=0, left_pt=0, right_pt=0):
    """设置表格单元格边距"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tblCellMar = OxmlElement('w:tblCellMar')

    if top_pt > 0:
        top = OxmlElement('w:top')
        top.set(qn('w:w'), str(int(top_pt * 20)))
        top.set(qn('w:type'), 'dxa')
        tblCellMar.append(top)

    if bottom_pt > 0:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:w'), str(int(bottom_pt * 20)))
        bottom.set(qn('w:type'), 'dxa')
        tblCellMar.append(bottom)

    if left_pt > 0:
        left = OxmlElement('w:left')
        left.set(qn('w:w'), str(int(left_pt * 20)))
        left.set(qn('w:type'), 'dxa')
        tblCellMar.append(left)

    if right_pt > 0:
        right = OxmlElement('w:right')
        right.set(qn('w:w'), str(int(right_pt * 20)))
        right.set(qn('w:type'), 'dxa')
        tblCellMar.append(right)

    # 移除旧的 tblCellMar
    old_mar = tblPr.find(qn('w:tblCellMar'))
    if old_mar is not None:
        tblPr.remove(old_mar)

    tblPr.append(tblCellMar)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_cell_double_border(cell, top=True, bottom=True):
    """设置单元格双线边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 移除旧边框
    old_borders = tcPr.find(qn('w:tcBorders'))
    if old_borders is not None:
        tcPr.remove(old_borders)

    tcBorders = OxmlElement('w:tcBorders')

    if top:
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'double')
        top_border.set(qn('w:sz'), '4')
        top_border.set(qn('w:color'), 'auto')
        tcBorders.append(top_border)

    if bottom:
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'double')
        bottom_border.set(qn('w:sz'), '4')
        bottom_border.set(qn('w:color'), 'auto')
        tcBorders.append(bottom_border)

    tcPr.append(tcBorders)


def set_cell_font(cell, font_name="宋体", font_size=12, bold=False, center=False):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold


def set_cell_text(cell, text, font_name="宋体", font_size=12, bold=False, center=False):
    """设置单元格文本和字体（简单文本）"""
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_font(cell, font_name, font_size, bold, center)


def set_cell_multiline_text(cell, text, font_name="宋体", font_size=12, bold=False):
    """设置单元格多行文本（使用硬回车/段落分隔）"""
    # 清空现有内容
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 按换行符分割，每行创建一个段落
    lines = text.split('\n') if text else [""]

    for i, line in enumerate(lines):
        if i == 0:
            # 第一行使用现有段落
            paragraph = cell.paragraphs[0]
        else:
            # 后续行添加新段落
            paragraph = cell.add_paragraph()

        run = paragraph.add_run(line.strip())
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold


def merge_cells(table, row_idx, start_col, end_col):
    """合并单元格"""
    cell = table.rows[row_idx].cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        cell.merge(table.rows[row_idx].cells[col])
    return cell


def set_column_widths(table, widths_cm):
    """设置列宽（厘米）"""
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def get_friday_of_week(year, week_num):
    """获取指定年份和周数的周五日期"""
    # 获取该年第一天
    jan1 = datetime(year, 1, 1)
    # 计算第一周的周一
    # ISO周：第一周是包含该年第一个周四的那周
    jan1_weekday = jan1.weekday()  # 0=周一, 6=周日

    if jan1_weekday <= 3:  # 周一到周四
        first_monday = jan1 - timedelta(days=jan1_weekday)
    else:  # 周五到周日
        first_monday = jan1 + timedelta(days=(7 - jan1_weekday))

    # 计算目标周的周一
    target_monday = first_monday + timedelta(weeks=week_num - 1)
    # 周五是周一+4天
    friday = target_monday + timedelta(days=4)

    return friday


def get_week_of_month(date):
    """获取日期是当月的第几周"""
    first_day = date.replace(day=1)
    # 计算第一天是星期几（周一=0）
    first_weekday = first_day.weekday()
    # 计算日期所在的周数
    adjusted_day = date.day + first_weekday
    week_num = (adjusted_day - 1) // 7 + 1
    return week_num


def generate_weekly_word(summary_data: dict, output_path: str = None) -> str:
    """生成周报Word文档 - 严格按照模版格式"""
    doc = Document()

    year = summary_data["year"]
    week_num = summary_data["week_num"]
    reports = summary_data.get("reports", [])

    # 计算周五日期和本月第几周
    friday = get_friday_of_week(year, week_num)
    month_week = get_week_of_month(friday)
    time_text = f"{friday.year}年{friday.month}月{friday.day}日，第{month_week}周工作周报"

    # 计算需要的行数
    this_week_rows = len(reports)
    next_week_rows = len(reports)

    # 总行数 = 2(标题+时间) + 1(本周表头) + 本周数据 + 1(空行) + 1(下周表头) + 下周数据
    total_rows = 2 + 1 + this_week_rows + 1 + 1 + next_week_rows

    # 创建表格：4列
    table = doc.add_table(rows=total_rows, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格单元格边距（上下5.65pt）
    set_table_cell_margins(table, top_pt=5.65, bottom_pt=5.65)

    # 设置列宽：1.74cm, 11.25cm, 2.00cm, 2.19cm（按模版调整）
    column_widths = [1.74, 11.25, 2.00, 2.19]
    set_column_widths(table, column_widths)

    # 第0行：产品研发部（合并所有单元格，行高33.25pt）
    set_row_height(table.rows[0], 33.25)
    merged_cell = merge_cells(table, 0, 0, 3)
    set_cell_text(merged_cell, "产品研发部", font_size=16, bold=True, center=True)

    # 第1行：时间信息（合并所有单元格，行高27.35pt）
    set_row_height(table.rows[1], 27.35)
    merged_cell = merge_cells(table, 1, 0, 3)
    set_cell_text(merged_cell, time_text, font_size=12, center=True)

    # 第2行：本周工作表头（行高25.25pt，双线边框）
    set_row_height(table.rows[2], 25.25)
    headers_this_week = ["项目", "本周工作说明", "执行人", "完成情况"]
    for i, header in enumerate(headers_this_week):
        cell = table.rows[2].cells[i]
        set_cell_text(cell, header, bold=True, center=True)
        set_cell_double_border(cell, top=True, bottom=True)

    # 第3行开始：本周工作数据
    data_start_row = 3
    for idx, report in enumerate(reports):
        row_idx = data_start_row + idx
        row = table.rows[row_idx].cells

        # 项目列留空
        set_cell_text(row[0], "", center=True)
        # 本周工作说明（使用硬回车）
        set_cell_multiline_text(row[1], report.get("this_week_work", "") or "")
        # 执行人
        set_cell_text(row[2], report.get("user_name", ""), center=True)
        # 完成情况
        set_cell_text(row[3], "完成", center=True)

    # 空行分隔（本周工作结束后）
    separator_row_idx = data_start_row + this_week_rows
    for cell in table.rows[separator_row_idx].cells:
        set_cell_text(cell, "")

    # 下周计划表头（行高28.9pt，双线边框）
    next_week_header_row = separator_row_idx + 1
    set_row_height(table.rows[next_week_header_row], 28.9)
    headers_next_week = ["项目", "下周工作计划", "执行人", ""]
    for i, header in enumerate(headers_next_week):
        cell = table.rows[next_week_header_row].cells[i]
        set_cell_text(cell, header, bold=True, center=True)
        set_cell_double_border(cell, top=True, bottom=True)

    # 下周计划数据
    next_week_data_start = next_week_header_row + 1
    for idx, report in enumerate(reports):
        row_idx = next_week_data_start + idx
        row = table.rows[row_idx].cells

        # 项目列留空
        set_cell_text(row[0], "", center=True)
        # 下周工作计划（使用硬回车）
        set_cell_multiline_text(row[1], report.get("next_week_plan", "") or "")
        # 执行人
        set_cell_text(row[2], report.get("user_name", ""), center=True)
        # 最后一列留空
        set_cell_text(row[3], "")

    # 保存文件
    if not output_path:
        os.makedirs(settings.DOCUMENTS_PATH, exist_ok=True)
        output_path = os.path.join(
            settings.DOCUMENTS_PATH,
            f"周报_{year}年第{week_num}周_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        )

    doc.save(output_path)
    return output_path
